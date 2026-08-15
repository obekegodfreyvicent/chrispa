#!/usr/bin/env python3
"""Converts a docs/*.md file to a matching *.docx, reproducing the styling
conventions already established in this repo's docs/ folder:
  - # / ## / ### -> Heading 1/2/3
  - **bold**, *italic*, `inline code` (Consolas 10pt) -> runs
  - [text](url) -> real hyperlink; relative ./NN-foo.md links retarget to .docx
  - - item / * item -> List Bullet (List Bullet 2/3 for indented items)
  - 1. item -> List Number (List Number 2/3 for indented items)
  - fenced ``` code blocks -> one Normal paragraph per line, Consolas 9.5pt
  - | a | b | tables -> Word table, style "Light Grid Accent 1", bold header row
  - ![alt](path) on its own line -> embedded picture (path resolved relative to
    the source .md's directory), 6" wide, with an italic caption paragraph
    below it using the alt text
  - horizontal rules (---) -> skipped
Blocks are emitted strictly in source order via document.add_*(), which
python-docx always appends at the true end of the body — so no low-level
XML repositioning is needed as long as we never go back and insert earlier.

Requires `python-docx` (`pip install python-docx`).

Usage:
  python3 md_to_docx.py <input.md> <output.docx>   # convert one file
  python3 md_to_docx.py --all                       # regenerate every
                                                      # docs/*.md -> matching
                                                      # .docx in this directory
"""
import os
import re
import sys

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

INLINE_CODE_FONT = 'Consolas'
INLINE_CODE_SIZE = Pt(10)
CODEBLOCK_FONT = 'Consolas'
CODEBLOCK_SIZE = Pt(9.5)
IMAGE_WIDTH = Inches(6)

IMAGE_RE = re.compile(r'^!\[(?P<alt>.*)\]\((?P<src>[^)]+)\)\s*$')

TOKEN_RE = re.compile(
    r'(?P<link>\[(?P<linktext>[^\]]+)\]\((?P<linkurl>[^)]+)\))'
    r'|(?P<code>`[^`]+`)'
    r'|(?P<bold>\*\*[^*]+\*\*)'
    r'|(?P<italic>\*[^*]+\*)'
)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(color)
    rPr.append(u)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def retarget_relative_md_link(url):
    if url.startswith('./') and '.md' in url:
        base = url.split('#', 1)[0]
        if base.endswith('.md'):
            return url.replace(base, base[:-3] + '.docx')
    return url


def is_special_line(line):
    stripped = line.strip()
    return (
        not stripped
        or line.startswith('#')
        or stripped.startswith('|')
        or re.match(r'^(\s*)[-*]\s+', line)
        or re.match(r'^(\s*)\d+\.\s+', line)
        or stripped.startswith('```')
        or re.match(r'^-{3,}$', stripped)
        or IMAGE_RE.match(stripped)
    )


def consume_continuations(lines, start, first_text):
    """Given the text of the line at `start`, absorb subsequent lines that
    are soft-wrap continuations (non-blank, not a new block) — this doc set
    hard-wraps prose/list-item text at ~100 cols with a 2-space-indented
    continuation rather than writing one physically long source line."""
    buf = [first_text]
    j = start + 1
    n = len(lines)
    while j < n and not is_special_line(lines[j]):
        buf.append(lines[j].strip())
        j += 1
    return ' '.join(buf), j


def apply_task_checkbox(text):
    text = re.sub(r'^\[ \]\s*', '☐ ', text)
    text = re.sub(r'^\[[xX]\]\s*', '☑ ', text)
    return text


def add_inline_runs(paragraph, text, bold=False, italic=False):
    """Recursive so bold/italic wrapping an inline-code span (e.g.
    **`SUPPORT_AGENT`**) still gets the Consolas code styling, not just the
    bold — ambient `bold`/`italic` carry down into nested matches."""
    pos = 0
    for m in TOKEN_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            if bold:
                run.bold = True
            if italic:
                run.italic = True
        if m.group('link'):
            label = m.group('linktext')
            url = retarget_relative_md_link(m.group('linkurl'))
            add_hyperlink(paragraph, url, label)
        elif m.group('code'):
            run = paragraph.add_run(m.group('code')[1:-1])
            run.font.name = INLINE_CODE_FONT
            run.font.size = INLINE_CODE_SIZE
            if bold:
                run.bold = True
            if italic:
                run.italic = True
        elif m.group('bold'):
            add_inline_runs(paragraph, m.group('bold')[2:-2], bold=True, italic=italic)
        elif m.group('italic'):
            add_inline_runs(paragraph, m.group('italic')[1:-1], bold=bold, italic=True)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        if bold:
            run.bold = True
        if italic:
            run.italic = True


def parse_table_rows(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith('|'):
        rows.append(lines[i].strip())
        i += 1

    def split_row(row):
        return [c.strip() for c in row.strip('|').split('|')]

    header = split_row(rows[0])
    data = [split_row(r) for r in rows[2:]] if len(rows) > 2 else []
    return header, data, i


def convert(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    d = docx.Document()
    title_text = None
    in_code = False
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('```'):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            p = d.add_paragraph()
            run = p.add_run(line if line else '')
            run.font.name = CODEBLOCK_FONT
            run.font.size = CODEBLOCK_SIZE
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if re.match(r'^-{3,}$', stripped):
            i += 1
            continue

        heading_m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if heading_m:
            level = len(heading_m.group(1))
            text = heading_m.group(2).strip()
            p = d.add_heading(level=level)
            add_inline_runs(p, text)
            if title_text is None and level == 1:
                title_text = text
            i += 1
            continue

        if stripped.startswith('|'):
            d.add_paragraph()
            header, data, i = parse_table_rows(lines, i)
            table = d.add_table(rows=1 + len(data), cols=len(header))
            table.style = 'Light Grid Accent 1'
            for c, htext in enumerate(header):
                cell_p = table.rows[0].cells[c].paragraphs[0]
                add_inline_runs(cell_p, htext)
                for r in cell_p.runs:
                    r.bold = True
            for r_idx, row in enumerate(data):
                for c_idx, ctext in enumerate(row):
                    if c_idx < len(header):
                        add_inline_runs(table.rows[r_idx + 1].cells[c_idx].paragraphs[0], ctext)
            continue

        list_m = re.match(r'^(\s*)[-*]\s+(.*)$', line)
        if list_m:
            indent = len(list_m.group(1))
            text, i = consume_continuations(lines, i, list_m.group(2))
            text = apply_task_checkbox(text)
            level = 1 + indent // 2
            style = 'List Bullet' if level <= 1 else f'List Bullet {min(level, 3)}'
            p = d.add_paragraph(style=style)
            add_inline_runs(p, text)
            continue

        num_m = re.match(r'^(\s*)\d+\.\s+(.*)$', line)
        if num_m:
            indent = len(num_m.group(1))
            text, i = consume_continuations(lines, i, num_m.group(2))
            level = 1 + indent // 2
            style = 'List Number' if level <= 1 else f'List Number {min(level, 3)}'
            p = d.add_paragraph(style=style)
            add_inline_runs(p, text)
            continue

        image_m = IMAGE_RE.match(stripped)
        if image_m:
            src = image_m.group('src')
            img_path = src if os.path.isabs(src) else os.path.join(os.path.dirname(os.path.abspath(md_path)), src)
            if os.path.isfile(img_path):
                d.add_picture(img_path, width=IMAGE_WIDTH)
                d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            alt = image_m.group('alt').strip()
            if alt:
                cap = d.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(alt)
                run.italic = True
                run.font.size = Pt(9.5)
            i += 1
            continue

        # plain paragraph: join consecutive non-blank, non-special lines
        text, i = consume_continuations(lines, i, line.strip())
        p = d.add_paragraph()
        add_inline_runs(p, text)

    if title_text:
        d.core_properties.title = title_text
    d.save(docx_path)


if __name__ == '__main__':
    if len(sys.argv) == 2 and sys.argv[1] == '--all':
        here = os.path.dirname(os.path.abspath(__file__))
        for name in sorted(os.listdir(here)):
            if name.endswith('.md'):
                md_path = os.path.join(here, name)
                docx_path = os.path.join(here, name[:-3] + '.docx')
                print(f'{name} -> {os.path.basename(docx_path)}')
                convert(md_path, docx_path)
    else:
        convert(sys.argv[1], sys.argv[2])
